import os
import tempfile
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF
from PIL import Image
import io

def convert_pdf_to_word(pdf_path, docx_path):
    """
    Convert PDF to Word document with better handling of tables and images
    """
    try:
        # Create a new Word document
        doc = Document()
        
        # Open the PDF
        pdf_document = fitz.open(pdf_path)
        
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            
            # Extract text
            text = page.get_text()
            if text.strip():
                # Add text to document
                paragraph = doc.add_paragraph()
                paragraph.add_run(text)
            
            # Extract images
            image_list = page.get_images(full=True)
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                
                # Convert to PIL Image
                image = Image.open(io.BytesIO(image_bytes))
                
                # Save temporarily
                temp_img_path = os.path.join(tempfile.gettempdir(), f"temp_img_{page_num}_{img_index}.png")
                image.save(temp_img_path)
                
                # Add image to document
                doc.add_picture(temp_img_path, width=Inches(6))
                
                # Clean up temp file
                os.remove(temp_img_path)
        
        # Save the document
        doc.save(docx_path)
        pdf_document.close()
        return True
        
    except Exception as e:
        print(f"Error converting PDF to Word: {e}")

        return False
